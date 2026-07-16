import streamlit as st
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from datetime import date

st.set_page_config(page_title="SBCI COD AUTOMATION", layout="centered")

st.markdown(
    """
    <div style="text-align: center; line-height: 1.2;">
        <h1 style="margin: 0; padding: 0;">SBC INSURANCE</h1>
        <p style="margin: 0; padding: 0; font-size: 35px;">
            Certificate of Destruction
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

# =========================
# CUSTOM INPUTS (WITH DEFAULTS + DROPDOWNS)
# =========================

col1, col2 = st.columns(2)

with col1:
    custom_date = st.date_input("📅 Certificate Date", value=date.today())
    prepared_date = st.date_input("📝 Prepared Date", value=date.today())

with col2:
    months_2026 = [
        "January 2026", "February 2026", "March 2026",
        "April 2026", "May 2026", "June 2026",
        "July 2026", "August 2026", "September 2026",
        "October 2026", "November 2026", "December 2026"
    ]

    custom_month = st.selectbox("📆 Month", months_2026)

    custom_reason = st.selectbox(
        "📌 Reason",
        ["Ceased/Paid", "End of Retention"]
    )

pull_out_date = st.date_input("📤 Pull-Out Date", value=date.today())

excel_file = st.file_uploader("📊 Upload Excel File", type=["xlsx"])
template_file = st.file_uploader("📄 Upload Word Template", type=["docx"])

# =========================
# SAFE TEXT REPLACEMENT (FIXED)
# =========================
def replace_text(doc, replacements, bold_keys=None):
    bold_keys = bold_keys or []

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:

                if key in bold_keys:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            run.bold = True
                else:
                    paragraph.text = paragraph.text.replace(key, value)

# =========================
# TABLE BORDERS
# =========================
def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)

# =========================
# MAIN TABLE
# =========================
def insert_table(doc, df):
    for paragraph in doc.paragraphs:
        if "{{TABLE}}" in paragraph.text:

            paragraph.text = ""

            table = doc.add_table(rows=1, cols=len(df.columns))
            table.autofit = True
            set_table_borders(table)

            for col_index, col_name in enumerate(df.columns):
                cell = table.rows[0].cells[col_index]
                cell.text = str(col_name)

                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True

                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), "4F81BD")
                tcPr.append(shd)

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for col_index, value in enumerate(row):
                    cell = row_cells[col_index]
                    cell.text = str(value) if pd.notna(value) else ""
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph._element.addnext(table._element)
            break

# =========================
# SUMMARY GENERATION
# =========================
def generate_summary(df):
    if "BUCKET LEVEL" not in df.columns or "BALANCE" not in df.columns:
        return None

    df["BALANCE"] = pd.to_numeric(df["BALANCE"], errors="coerce").fillna(0)

    summary = df.groupby("BUCKET LEVEL").agg(
        TOTAL_ACCTS=("BUCKET LEVEL", "count"),
        TOTAL_BALANCE=("BALANCE", "sum")
    ).reset_index()

    total_accts = summary["TOTAL_ACCTS"].sum()
    total_balance = summary["TOTAL_BALANCE"].sum()

    return summary, total_accts, total_balance

# =========================
# SUMMARY TABLE
# =========================
def insert_summary_table(doc, summary, total_accts, total_balance):
    for paragraph in doc.paragraphs:
        if "{{SUMMARY_TABLE}}" in paragraph.text:

            paragraph.text = ""

            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            set_table_borders(table)

            summary_row = table.rows[0].cells
            merged_cell = summary_row[0].merge(summary_row[1]).merge(summary_row[2])
            merged_cell.text = "SUMMARY"

            for p in merged_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True

            tc = merged_cell._element
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), "4F81BD")
            tcPr.append(shd)

            headers = ["Bucket Level", "Total # of Accts", "Total Balance"]
            header_cells = table.add_row().cells

            for i, header in enumerate(headers):
                cell = header_cells[i]
                cell.text = header

                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True

            for _, row in summary.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row["BUCKET LEVEL"])
                row_cells[1].text = str(row["TOTAL_ACCTS"])
                row_cells[2].text = f"{row['TOTAL_BALANCE']:,.2f}"

                for cell in row_cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            total_row = table.add_row().cells
            total_row[0].text = "Grand Total"
            total_row[1].text = str(total_accts)
            total_row[2].text = f"{total_balance:,.2f}"

            for cell in total_row:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True

            paragraph._element.addnext(table._element)
            break

# =========================
# PRODUCT + BUCKET TEXT
# =========================
def generate_product_bucket_text(df):
    if "PRODUCT" not in df.columns or "BUCKET LEVEL" not in df.columns:
        return ""

    products = sorted(df["PRODUCT"].dropna().astype(str).str.strip().unique())
    buckets = sorted(df["BUCKET LEVEL"].dropna().astype(str).str.strip().unique())

    return f"<{', '.join(products)} / {', '.join(buckets)}>"

# =========================
# MAIN PROCESS
# =========================
if excel_file and template_file:
    try:
        df = pd.read_excel(excel_file)

        df.columns = df.columns.str.strip().str.upper()
        original_df = df.copy()

        rename_map = {}
        for col in df.columns:
            if "IS" in col:
                rename_map[col] = "IS#"
            if "ENDO" in col:
                rename_map[col] = "ENDO DATE"
            if "PULL" in col:
                rename_map[col] = "PULL-OUT DATE"

        df = df.rename(columns=rename_map)

        if "NO" not in df.columns:
            df.insert(0, "NO", range(1, len(df) + 1))

        formatted_date = custom_date.strftime("%B %d, %Y")
        formatted_prepared_date = prepared_date.strftime("%m/%d/%Y")
        formatted_pullout_date = pull_out_date.strftime("%m/%d/%Y")

        df["PULL-OUT DATE"] = formatted_pullout_date

        if "ENDO DATE" in df.columns:
            df["ENDO DATE"] = pd.to_datetime(df["ENDO DATE"], errors="coerce").dt.strftime("%m/%d/%Y")

        df = df[["NO", "IS#", "ENDO DATE", "PULL-OUT DATE"]]

        st.success("✅ Excel Loaded Successfully!")
        st.dataframe(df)

        product_bucket_text = generate_product_bucket_text(original_df)

        doc = Document(template_file)

        replacements = {
            "{{DATE}}": formatted_date,
            "{{MONTH}}": custom_month,
            "{{REASON}}": custom_reason,
            "{{PREPARED_DATE}}": formatted_prepared_date,
            "{{PRODUCT_BUCKET}}": product_bucket_text
        }

        replace_text(
            doc,
            replacements,
            bold_keys=["{{MONTH}}", "{{DATE}}"]
        )

        insert_table(doc, df)

        summary_data = generate_summary(original_df)

        if summary_data:
            summary_df, total_accts, total_balance = summary_data
            insert_summary_table(doc, summary_df, total_accts, total_balance)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("🎉 Certificate Generated Successfully!")

        st.download_button(
            label="📥 Download Certificate Of Destruction",
            data=buffer,
            file_name="SBCI_Certificate of Destruction.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"❌ Error: {e}")
