import streamlit as st
import pandas as pd
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
from copy import copy
import os
import re
from datetime import date, datetime
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(layout="wide")

if "current_page" not in st.session_state:
    st.session_state.current_page = "daily_report"

# css
st.markdown(
    """
    <style>
    /* background */
    .stApp {
        background-color: #FFF0F5 !important; 
    }
    
    [data-testid="stHeader"] {
        background-color: rgba(0, 0, 0, 0) !important;
        background-image: none !important;
        border: none !important;
    }
    
    /* sidebar arrow button */
    [data-testid="stHeader"] button, 
    button[aria-label="Close sidebar"], 
    button[aria-label="Open sidebar"] {
        color: #FF1493 !important; 
        background-color: #FFE4E1 !important; 
        border: 1px solid #FFB6C1 !important;
        border-radius: 50% !important;
        padding: 4px !important;
        transition: all 0.3s ease !important;
    }
    [data-testid="stHeader"] button:hover {
        background-color: #FF69B4 !important;
        color: #FFFFFF !important;
    }
    
    [data-testid="stAppViewContainer"] {
        top: 0px;
    }

    /* sidebar */
    section[data-testid="stSidebar"] {
        background-color: #FFE4E1 !important; 
        border-right: 1px solid #FFD1DC !important;
    }
    
    .stApp, .stApp p, .stApp span, .stApp label, .stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp h5, .stApp h6,
    section[data-testid="stSidebar"] p, section[data-testid="stSidebar"] span, section[data-testid="stSidebar"] label {
        color: #1A1A1A !important;
    }
    
    .stApp small, .stApp .st-emotion-cache-183lyg1, .stApp [data-testid="stMarkdownContainer"] p {
        color: #2D2D2D !important;
    }

    section[data-testid="stSidebar"] [data-testid="stNotification"] {
        background-color: #D1E7DD !important; 
        border: 1px solid #BADBCC !important;
        border-radius: 8px !important;
        text-align: center !important;
        display: flex !important;
        justify-content: center !important;
        align-items: center !important;
        padding: 8px 12px !important;
    }
    section[data-testid="stSidebar"] [data-testid="stNotification"] p,
    section[data-testid="stSidebar"] [data-testid="stNotification"] span {
        color: #0F5132 !important; 
        font-size: 14px !important;
        font-weight: normal !important;
    }
    section[data-testid="stSidebar"] [data-testid="stNotification"] svg {
        color: #0F5132 !important; 
        fill: #0F5132 !important;
    }
    section[data-testid="stSidebar"] [data-testid="stNotification"] [data-testid="stMarkdownContainer"] {
        width: 100% !important;
        text-align: center !important;
    }

    .stApp [data-testid="stMain"] [data-testid="stNotification"] {
        color: #0F5132 !important;
    }
    .stApp [data-testid="stMain"] [data-testid="stNotification"] p {
        color: #0F5132 !important;
    }
    
    /* file uploader */
    div[data-testid="stFileUploader"] section {
        background-color: #FFF5F7 !important; 
        border: 2px dashed #FFB6C1 !important;
        border-radius: 12px !important;
        padding: 1.5rem !important;
    }
    
    div[data-testid="stFileUploader"] button {
        background-color: #FFC0CB !important; 
        color: #1A1A1A !important;
        font-weight: bold !important;
        border: 1px solid #FFB6C1 !important;
        border-radius: 8px !important;
        transition: all 0.3s ease-in-out !important;
    }
    
    div[data-testid="stFileUploader"] button:hover {
        background-color: #FF69B4 !important; 
        color: #FFFFFF !important;
        border-color: #FF1493 !important;
    }
    
    [data-testid="stFileUploaderFileName"],
    [data-testid="stFileUploaderFile"] span,
    [data-testid="stFileUploaderFile"] div,
    [data-testid="stFileUploaderFileName"] p {
        color: #2D2D2D !important; 
    }
    [data-testid="stFileUploaderFile"] {
        background-color: #FFE4E1 !important;
        border: 1px solid #FFC0CB !important;
        border-radius: 8px !important;
    }
    
    [data-testid="stDataFrame"] {
        background-color: #FFFFFF !important;
        border: 1px solid #FFD1DC !important;
        border-radius: 8px !important;
    }
    [data-testid="stDataFrame-toolbar"] {
        background-color: #FFE4E1 !important;
        border-radius: 4px !important;
    }
    [data-testid="stDataFrame-toolbar"] button {
        color: #FF1493 !important; 
        fill: #FF1493 !important;
    }
    [data-testid="stDataFrame-toolbar"] button:hover {
        background-color: #FFC0CB !important;
    }
    
    div.stButton > button {
        background-color: #FFC0CB !important; 
        color: #1A1A1A !important;
        font-weight: bold !important;
        border: 1px solid #FFB6C1 !important;
        border-radius: 8px !important;
        padding: 0.5rem 1rem !important;
        transition: all 0.3s ease-in-out !important;
    }
    
    div.stButton > button:hover {
        background-color: #FF69B4 !important; 
        color: #FFFFFF !important;
        border-color: #FF1493 !important;
        transform: scale(1.02);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
    }
    
    div.stDownloadButton > button {
        background-color: #FFB6C1 !important;
        color: #1A1A1A !important;
        font-weight: bold !important;
        border: 1px solid #FFA07A !important;
        border-radius: 8px !important;
        transition: all 0.3s ease-in-out !important;
    }
    div.stDownloadButton > button:hover {
        background-color: #FF69B4 !important; 
        color: #FFFFFF !important;
        transform: scale(1.02);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
    }

    div.stButton > button[key="green_nav"] {
        background-color: #2E7D32 !important; 
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        padding: 15px 10px !important;
        font-size: 15px !important;
        font-weight: bold !important;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1) !important;
        transition: all 0.3s ease !important;
    }
    div.stButton > button[key="green_nav"]:hover {
        background-color: #1B5E20 !important; 
        color: #FFFFFF !important;
        transform: scale(1.02);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15) !important;
    }

    div.stButton > button[key="blue_nav"] {
        background-color: #1565C0 !important; 
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        padding: 15px 10px !important;
        font-size: 15px !important;
        font-weight: bold !important;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1) !important;
        transition: all 0.3s ease !important;
    }
    div.stButton > button[key="blue_nav"]:hover {
        background-color: #0D47A1 !important; 
        color: #FFFFFF !important;
        transform: scale(1.02);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15) !important;
    }

    div.stButton > button[key="cod_nav"] {
        background-color: #FFC0CB !important; 
        color: #1A1A1A !important;
        font-weight: bold !important;
        border: 1px solid #FFB6C1 !important;
        border-radius: 8px !important;
        padding: 12px 10px !important;
        font-size: 15px !important;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05) !important;
        transition: all 0.3s ease-in-out !important;
    }
    div.stButton > button[key="cod_nav"]:hover {
        background-color: #FF69B4 !important; 
        color: #FFFFFF !important;
        border-color: #FF1493 !important;
        transform: scale(1.02);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15) !important;
    }

    div.stButton > button[key="reset_template_btn"] {
        width: 100% !important;
        text-align: center !important;
        display: block !important;
        padding: 12px 10px !important;
        font-size: 14px !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

TEMPLATE_PATH = r"C:\Users\SPM\Downloads\EFS DAR TEMPLATE______.xlsx"

# date format
def generate_month_year_suffix(df):
    if 'Action Date' in df.columns and not df['Action Date'].dropna().empty:
        try:
            dates = pd.to_datetime(df['Action Date'].dropna(), errors='coerce').dropna()
            if not dates.empty:
                sample_date = dates.iloc[0]
                return f"for {sample_date.strftime('%B %Y').upper()}"
        except:
            pass
    return f"for {pd.Timestamp.now().strftime('%B %Y').upper()}"


# sidebar
with st.sidebar:
    with st.expander("🧭 EFS-DAR Menu", expanded=False):
        # menu 1
        if st.button("Daily Remark Report", use_container_width=True):
            st.session_state.current_page = "daily_report"
            st.rerun()
            
        # menu 2
        if st.button("Merge EFS DAR Files", key="green_nav", use_container_width=True):
            st.session_state.current_page = "merge_page"
            st.rerun()

        # menu 3
        if st.button("RFD Summary", key="blue_nav", use_container_width=True):
            st.session_state.current_page = "rfd_summary_page"
            st.rerun()

    st.markdown("---")

    # cod automation
    if st.button("COD Automation", key="cod_nav", use_container_width=True):
        st.session_state.current_page = "cod_automation"
        st.rerun()

    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown("---")

    st.caption("⚙️ **System Configurations**")
    if os.path.exists(TEMPLATE_PATH):
        st.success("Network Connected.")
        if st.button("Change Master Template", key="reset_template_btn"):
            try:
                os.remove(TEMPLATE_PATH)
                st.rerun()
            except Exception as e:
                st.error(f"Could not delete network file automatically: {e}")


# UNANG MEOWNU PO ITO PARA SA DAR ===============================================================================================================================================================
if st.session_state.current_page == "daily_report":

    st.title("Daily Remark Report | SBC")

    template_exists = os.path.exists(TEMPLATE_PATH)

    if not template_exists:
        st.subheader("EFS DAR TEMPLATE")
        template_file = st.file_uploader("Template path not found on network share. Upload your Excel template manually", type=["xlsx"])
        
        if template_file is not None:
            try:
                with open(TEMPLATE_PATH, "wb") as f:
                    f.write(template_file.read())
                st.success("Template saved permanently to your network shared folder!")
                st.rerun()
            except Exception as e:
                st.error(f"Could not save directly to network path due to system permissions. Hardcoded path error: {e}")

    else:
        uploaded_file = st.file_uploader("Choose Daily Remark Report Excel file", type=["xlsx", "xls"])

        if uploaded_file is not None:
            try:
                df_source = pd.read_excel(uploaded_file)
                st.success("Daily report uploaded successfully!")
                
                if not df_source.empty:
                    last_row_idx = df_source.index[-1]
                    first_cell = df_source.iloc[-1, 0]
                    other_cells = df_source.iloc[-1, 1:]
                    
                    if pd.notna(first_cell) and str(first_cell).strip() != "":
                        if other_cells.isna().all() or other_cells.apply(lambda x: str(x).strip() == "").all():
                            df_source = df_source.drop(last_row_idx)
                
                # date
                extracted_date_str = "REPORT_DATE"
                if 'Date' in df_source.columns and not df_source['Date'].dropna().empty:
                    first_date = pd.to_datetime(df_source['Date'].dropna().iloc[0], errors='coerce')
                    if pd.notna(first_date):
                        extracted_date_str = first_date.strftime('%m/%d/%Y')
                    
                template_columns = [
                    'Action Date', 'Campaign', 'Agency', 'IS #', 'Account', 
                    'Action', 'Result', 'Contact', 'Payment_Date', 'Payment_Amount', 
                    'RFD', 'Comments', 'COUNTER', 'Product', 'Level', 'Remarks'
                ]
                
                df_output = pd.DataFrame(columns=template_columns)
                
                # mapping
                if 'Date' in df_source.columns:
                    df_output['Action Date'] = pd.to_datetime(df_source['Date'], errors='coerce').dt.strftime('%m/%d/%Y')
                if 'Account No.' in df_source.columns:
                    df_output['IS #'] = df_source['Account No.']
                if 'Card No.' in df_source.columns:
                    df_output['Account'] = df_source['Card No.']
                if 'Remark Type' in df_source.columns:
                    df_output['Action'] = df_source['Remark Type']
                if 'Status' in df_source.columns:
                    df_output['Result'] = df_source['Status']
                if 'Remark' in df_source.columns:
                    df_output['Remarks'] = df_source['Remark']
                if 'Product Type' in df_source.columns:
                    df_output['Product'] = df_source['Product Type']
                    
                df_output['Campaign'] = "ACA"
                df_output['Agency'] = "SP Madrid"
                
                df_output['Product'] = df_output['Product'].fillna("").astype(str).str.strip()
                df_output['Remarks'] = df_output['Remarks'].fillna("").astype(str).str.strip()
                df_output['Result'] = df_output['Result'].fillna("").astype(str).str.strip()
                
                # formula para sa column L
                df_output['Comments'] = df_output.apply(
                    lambda r: f"{r['Product']} | {r['Remarks']}" if r['Product'] or r['Remarks'] else "", 
                    axis=1
                )
                
                # ito naman sa column M
                df_output['COUNTER'] = df_output['Comments'].apply(len)

                # mga rfd na sana hindi ma-overwrite ever pls pls
                whitelist_reasons = [
                    "TFIP", "BUSY/OVERLOOKED", "BUSINESS DOWN/DELAYED COLLECTIONS", "OOT/OOC",
                    "MEDICAL REASON/PRIORITIZED MEDICAL AND HOSPITAL EXPENSES", "UNEMPLOYED",
                    "BANK PROCESS ISSUE", "DISPUTE", "SQUANDERED BY 3RD PARTY", "CALAMITY",
                    "DECEASED", "FRAUD", "NSOA/LSOA", "NIOP", "FREE TEXT"
                ]

                def process_rfd_routing(comments_val, result_val):
                    comments_upper = str(comments_val).strip().upper()
                    result_upper = str(result_val).strip().upper()

                    if not comments_upper and not result_upper:
                        return "UNCONTACTABLE"

                    combined = f"{comments_upper} {result_upper}"
                    
                    if "DEAR" in combined or "####" in combined or "UNCONTACT" in combined or combined.strip() == "":
                        return "UNCONTACTABLE"
                    if "BUSY" in combined or "OVERLOOKED" in combined:
                        return "BUSY/OVERLOOKED"
                    if "MEDICAL" in combined or "HOSPITAL" in combined:
                        return "MEDICAL REASON/PRIORITIZED MEDICAL AND HOSPITAL EXPENSES"
                    if "UNEMPLOY" in combined:
                        return "UNEMPLOYED"
                    if "DISPUTE" in combined:
                        return "DISPUTE"
                    if "FRAUD" in combined:
                        return "FRAUD"
                    if "DECEASED" in combined or "DEAD" in combined:
                        return "DECEASED"
                    if "CALAMITY" in combined:
                        return "CALAMITY"

                    extracted_bracket_val = ""
                    brackets = re.findall(r'\[([^\]]+)\]', comments_val)
                    if brackets:
                        extracted_bracket_val = brackets[-1].strip().upper()

                    if extracted_bracket_val:
                        for reason in whitelist_reasons:
                            clean_reason = re.sub(r'[^A-Z0-9]', '', reason)
                            clean_extracted = re.sub(r'[^A-Z0-9]', '', extracted_bracket_val)
                            if clean_reason == clean_extracted or clean_extracted in clean_reason:
                                return reason
                        return "UNCONTACTABLE"
                    return "UNCONTACTABLE"
            
                df_output['RFD'] = df_output.apply(
                    lambda r: process_rfd_routing(r['Comments'], r['Result']), 
                    axis=1
                )
                

                for col in template_columns:
                    if col not in df_output.columns:
                        df_output[col] = ""
                    else:
                        df_output[col] = df_output[col].fillna("")

                st.subheader(" Automated Template Output Preview")
                st.dataframe(df_output[template_columns].head(5), height=215, use_container_width=True)
                
                output_buffer = BytesIO()
                wb = openpyxl.load_workbook(TEMPLATE_PATH)
                ws = wb.active
                
                template_styles = {}
                for col_idx in range(1, len(template_columns) + 1):
                    sample_cell = ws.cell(row=3, column=col_idx)
                    template_styles[col_idx] = {
                        'font': copy(sample_cell.font),
                        'border': copy(sample_cell.border),
                        'fill': copy(sample_cell.fill),
                        'alignment': copy(sample_cell.alignment),
                        'number_format': sample_cell.number_format
                    }

                current_write_row = 4
                for index, row in df_output.iterrows():
                    for col_idx, col_name in enumerate(template_columns, start=1):
                        target_cell = ws.cell(row=current_write_row, column=col_idx)
                        target_cell.value = row[col_name]
                        
                        style = template_styles.get(col_idx)
                        if style:
                            if style['font']: target_cell.font = style['font']
                            if style['border']: target_cell.border = style['border']
                            if style['fill']: target_cell.fill = style['fill']
                            if style['alignment']: target_cell.alignment = style['alignment']
                            target_cell.number_format = style['number_format']
                    
                    current_write_row += 1

                ws.delete_rows(3, amount=1)
                
                header_font = Font(name="Tahoma", size=10, bold=True)
                for r in [1, 2]:
                    for c in range(1, len(template_columns) + 1):
                        cell = ws.cell(row=r, column=c)
                        if cell.value:
                            cell.font = header_font
                
                wb.save(output_buffer)
                processed_bytes = output_buffer.getvalue()
                
                safe_filename_date = extracted_date_str.replace('/', '-')
                final_filename = f"EFS DAR {safe_filename_date}.xlsx"
                
                st.subheader("Your file is ready.")
                st.download_button(
                    label=f"Download {final_filename}",
                    data=processed_bytes,
                    file_name=final_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            except Exception as e:
                st.error(f"An error occurred while compiling your data layout: {e}")

# ITO NA PO YUNG MEOWNU MERGER ===============================================================================================================================================================
elif st.session_state.current_page == "merge_page":
    st.title("Merge EFS DAR Files")

    uploaded_files = st.file_uploader(
        "Choose EFS DAR Excel files to merge", 
        type=["xlsx"], 
        accept_multiple_files=True
    )

    if uploaded_files:
        st.success(f"Loaded {len(uploaded_files)} files successfully!")
        
        if st.button("Combine & Process Files", type="primary"):
            all_dfs = []
            
            for file in uploaded_files:
                try:
                    df = pd.read_excel(file)
                    all_dfs.append(df)
                except Exception as e:
                    st.error(f"Error reading file {file.name}: {e}")
            
            if all_dfs:
                merged_df = pd.concat(all_dfs, ignore_index=True)
                
                if 'Action Date' in merged_df.columns:
                    merged_df['Action Date'] = merged_df['Action Date'].astype(str).str.strip()
                    
                    merged_df = merged_df[~merged_df['Action Date'].str.upper().isin(['ACTION DATE', 'MM/DD/YYYY', 'MM-DD-YYYY', 'DD/MM/YYYY', 'YYYY-MM-DD'])]
                    date_template_pattern = r'(MM|DD|YYYY)'
                    merged_df = merged_df[~merged_df['Action Date'].str.contains(date_template_pattern, case=False, na=False, regex=True)]
                    
                    parsed_dates = pd.to_datetime(merged_df['Action Date'], errors='coerce')
                    merged_df['temp_sort_date'] = parsed_dates
                    
                    merged_df = merged_df.dropna(subset=['temp_sort_date']).reset_index(drop=True)
                    
                    merged_df = merged_df.sort_values(by='temp_sort_date').reset_index(drop=True)
                    merged_df['Action Date'] = merged_df['temp_sort_date'].dt.strftime('%m/%d/%Y')
                    merged_df = merged_df.drop(columns=['temp_sort_date'])

                # very forgiving
                if 'RFD' in merged_df.columns:
                    def clean_merged_rfd(val):
                        if pd.isna(val):
                            return val
                        val_str = str(val).strip().upper()

                        if "DIPUTE" in val_str or "DISPUT" in val_str:
                            return "DISPUTE"
                        if "UNCONTACT" in val_str or "DEAR" in val_str or "####" in val_str:
                            return "UNCONTACTABLE"
                        if "BUSY" in val_str or "OVERLOOK" in val_str:
                            return "BUSY/OVERLOOKED"
                        if "MEDICAL" in val_str or "HOSPITAL" in val_str:
                            return "MEDICAL REASON/PRIORITIZED MEDICAL AND HOSPITAL EXPENSES"
                        if "UNEMPLOY" in val_str:
                            return "UNEMPLOYED"
                        if "DECEASE" in val_str or "DEAD" in val_str:
                            return "DECEASED"
                        if "BANK PROCESS" in val_str:
                            return "BANK PROCESS ISSUE"
                        if "BUSINESS DOWN" in val_str or "DELAYED COLL" in val_str:
                            return "BUSINESS DOWN/DELAYED COLLECTIONS"
                        if "3RD PARTY" in val_str or "SQUANDER" in val_str:
                            return "SQUANDERED BY 3RD PARTY"
                        if "FRAUD" in val_str:
                            return "FRAUD"
                        if "CALAMITY" in val_str:
                            return "CALAMITY"

                        return val_str

                    merged_df['RFD'] = merged_df['RFD'].apply(clean_merged_rfd)
                
                st.success("Files successfully combined into a single dataset!")
                st.subheader("Combined File Preview")
                st.dataframe(merged_df, use_container_width=True)
                
                output_buffer = BytesIO()
                
                if os.path.exists(TEMPLATE_PATH):
                    wb = openpyxl.load_workbook(TEMPLATE_PATH)
                    ws = wb.active
                    
                    template_styles = {}
                    columns_count = merged_df.shape[1]
                    for col_idx in range(1, columns_count + 1):
                        sample_cell = ws.cell(row=3, column=col_idx)
                        template_styles[col_idx] = {
                            'font': copy(sample_cell.font),
                            'border': copy(sample_cell.border),
                            'fill': copy(sample_cell.fill),
                            'alignment': copy(sample_cell.alignment),
                            'number_format': sample_cell.number_format
                        }

                    current_write_row = 4
                    for index, row in merged_df.iterrows():
                        for col_idx in range(1, columns_count + 1):
                            target_cell = ws.cell(row=current_write_row, column=col_idx)
                            target_cell.value = row.iloc[col_idx - 1]
                            
                            style = template_styles.get(col_idx)
                            if style:
                                if style['font']: target_cell.font = style['font']
                                if style['border']: target_cell.border = style['border']
                                if style['fill']: target_cell.fill = style['fill']
                                if style['alignment']: target_cell.alignment = style['alignment']
                                target_cell.number_format = style['number_format']
                        current_write_row += 1

                    # row guide
                    ws.delete_rows(3, amount=1)
                    
                    header_font = Font(name="Tahoma", size=10, bold=True)
                    for r in [1, 2]:
                        for c in range(1, columns_count + 1):
                            cell = ws.cell(row=r, column=c)
                            if cell.value:
                                cell.font = header_font
                    
                    wb.save(output_buffer)
                else:
                    with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
                        merged_df.to_excel(writer, index=False, sheet_name='Merged DAR')
                
                month_year_suffix = generate_month_year_suffix(merged_df)
                merged_filename = f"DAR {month_year_suffix}.xlsx"

                st.subheader("Your file is ready:")
                st.download_button(
                    label=f"Download {merged_filename}",
                    data=output_buffer.getvalue(),
                    file_name=merged_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )


# RFD SUMMARY CALCULATOR KUNO ===============================================================================================================================================================
elif st.session_state.current_page == "rfd_summary_page":
    st.title("RFD Summary Calculator")

    uploaded_summary_file = st.file_uploader("Upload Merged EFS DAR Excel File", type=["xlsx"])

    if uploaded_summary_file is not None:
        try:
            wb_check = openpyxl.load_workbook(uploaded_summary_file)
            sheet_name = wb_check.sheetnames[0]
            ws_check = wb_check[sheet_name]
            
            df_summary = pd.read_excel(uploaded_summary_file)

            # delete yung second header
            if not df_summary.empty:
                df_summary = df_summary.drop(df_summary.index[0]).reset_index(drop=True)

            # delete mga empty columns
            cols_to_drop = []
            for idx in [7, 8, 9, 14]: # H, I, J, O
                if idx < len(df_summary.columns):
                    cols_to_drop.append(df_summary.columns[idx])
            
            df_cleaned = df_summary.drop(columns=cols_to_drop, errors='ignore')

            key_cols = ['IS #', 'RFD', 'Product']
            missing_cols = [c for c in key_cols if c not in df_cleaned.columns]
            
            if missing_cols:
                st.error(f"Missing required columns in the file: {missing_cols}. Please ensure the file has 'IS #', 'RFD', and 'Product' headers.")
            else:
                # hierarchy sa RFD
                priority_list = [
                    "DISPUTE",
                    "BANK PROCESS ISSUE",
                    "DECEASED",
                    "MEDICAL REASON/PRIORITIZED MEDICAL AND HOSPITAL EXPENSES",
                    "MEDICAL REASON/PRIORITIZED MEDICAL AND HOSPITAL EXPENSE", # handling singular/plural spelling
                    "CALAMITY",
                    "BUSINESS DOWN/DELAYED COLLECTIONS",
                    "TFIP",
                    "BUSY/OVERLOOKED",
                    "UNEMPLOYED",
                    "OOT/OOC",
                    "SQUANDERED BY 3RD PARTY",
                    "FRAUD",
                    "NSOA/LSOA",
                    "NIOP",
                    "UNCONTACTABLE"
                ]
                
                df_cleaned['RFD_clean'] = df_cleaned['RFD'].astype(str).str.strip().str.upper()
                df_cleaned['Product_clean'] = df_cleaned['Product'].astype(str).str.strip().str.upper()
                
                priority_map = {rfd.upper(): idx for idx, rfd in enumerate(priority_list)}
                
                df_cleaned['Priority_Rank'] = df_cleaned['RFD_clean'].map(priority_map).fillna(99)
                
                # sort by unique and rfd
                df_sorted = df_cleaned.sort_values(by=['IS #', 'Priority_Rank'], ascending=[True, True])
                
                # keep lang yung highest kase ba naman ano naman gagawin natin diyan sa deceased and uncontactable, di ba? so keep first lang
                df_deduped = df_sorted.drop_duplicates(subset=['IS #'], keep='first').copy()
                
                st.success(f"File uploaded successfully! Unique IS count: {len(df_deduped)} (from {len(df_cleaned)} raw entries)")

                # summarizination
                allowed_products = ['AL', 'SBL', 'SHL']
                df_matrix_source = df_deduped[df_deduped['Product_clean'].isin(allowed_products)].copy()
                
                matrix = pd.crosstab(
                    df_matrix_source['RFD_clean'], 
                    df_matrix_source['Product_clean'], 
                    margins=False
                )
                
                for prod in allowed_products:
                    if prod not in matrix.columns:
                        matrix[prod] = 0
                
                # AL - SBL - SHL
                matrix = matrix[allowed_products]
                
                existing_rfds = [p.upper() for p in priority_list if p.upper() in matrix.index]
                other_rfds = [idx for idx in matrix.index if idx not in existing_rfds]
                matrix = matrix.reindex(existing_rfds + other_rfds).fillna(0).astype(int)
                
                # sort 
                matrix['Grand Total'] = matrix.sum(axis=1)
                matrix = matrix.sort_values(by='Grand Total', ascending=False)
                
                grand_totals = matrix.sum(axis=0)
                matrix.loc['Grand Total'] = grand_totals
                
                matrix = matrix.reset_index().rename(columns={'RFD_clean': 'RFD'})

                st.subheader("Calculated Unique RFD Summary Table")
                st.dataframe(matrix, use_container_width=True)

                # excel format
                df_final_cleaned = df_deduped.drop(columns=['RFD_clean', 'Product_clean', 'Priority_Rank'], errors='ignore')
                output_buffer = BytesIO()

                if os.path.exists(TEMPLATE_PATH):
                    wb_styled = openpyxl.load_workbook(TEMPLATE_PATH)
                    ws_dar = wb_styled.active
                    ws_dar.title = 'DAR'
                    
                    sorted_drops = sorted([7, 8, 9, 14], reverse=True)
                    for col_idx_zero_based in sorted_drops:
                        if col_idx_zero_based < ws_dar.max_column:
                            ws_dar.delete_cols(col_idx_zero_based + 1)
                    
                    template_styles = {}
                    columns_count = df_final_cleaned.shape[1]
                    for col_idx in range(1, columns_count + 1):
                        sample_cell = ws_dar.cell(row=3, column=col_idx)
                        template_styles[col_idx] = {
                            'font': copy(sample_cell.font),
                            'border': copy(sample_cell.border),
                            'fill': copy(sample_cell.fill),
                            'alignment': copy(sample_cell.alignment),
                            'number_format': sample_cell.number_format
                        }
                        
                    current_write_row = 4
                    for index, row in df_final_cleaned.iterrows():
                        for col_idx in range(1, columns_count + 1):
                            target_cell = ws_dar.cell(row=current_write_row, column=col_idx)
                            target_cell.value = row.iloc[col_idx - 1]
                            
                            style = template_styles.get(col_idx)
                            if style:
                                if style['font']: target_cell.font = style['font']
                                if style['border']: target_cell.border = style['border']
                                if style['fill']: target_cell.fill = style['fill']
                                if style['alignment']: target_cell.alignment = style['alignment']
                                target_cell.number_format = style['number_format']
                        current_write_row += 1

                    # row guide
                    ws_dar.delete_rows(3, amount=1)
                    
                    header_font = Font(name="Tahoma", size=10, bold=True)
                    for r in [1, 2]:
                        for c in range(1, columns_count + 1):
                            cell = ws_dar.cell(row=r, column=c)
                            if cell.value:
                                cell.font = header_font
                                
                    # RFD SUMMARY TABLE: sa second sheet ====================================================================================
                    ws_summary = wb_styled.create_sheet(title='RFD Summary Matrix')
                    ws_summary.views.sheetView[0].showGridLines = True
                    
                    report_month_year = "SUMMARY REPORT"
                    if 'Action Date' in df_final_cleaned.columns and not df_final_cleaned['Action Date'].dropna().empty:
                        try:
                            sample_date = pd.to_datetime(df_final_cleaned['Action Date'].dropna().iloc[0], errors='coerce')
                            if pd.notna(sample_date):
                                report_month_year = sample_date.strftime('%B %Y').upper()
                        except:
                            pass
                    
                    matrix_styled = matrix.rename(columns={'Grand Total': 'TOTAL'})
                    
                    # format
                    font_family = "Tahoma"
                    green_fill = PatternFill(start_color="2E7D32", end_color="2E7D32", fill_type="solid")
                    mint_header_fill = PatternFill(start_color="66BB6A", end_color="66BB6A", fill_type="solid")
                    zebra_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
                    
                    font_title = Font(name=font_family, size=14, bold=True, color="FFFFFF")
                    font_header = Font(name=font_family, size=10, bold=True, color="FFFFFF")
                    font_body_bold = Font(name=font_family, size=10, bold=True)
                    font_body_regular = Font(name=font_family, size=10)
                    
                    thin_border = Border(
                        left=Side(style='thin', color='CCCCCC'),
                        right=Side(style='thin', color='CCCCCC'),
                        top=Side(style='thin', color='CCCCCC'),
                        bottom=Side(style='thin', color='CCCCCC')
                    )
                    
                    # month of the report
                    num_cols = matrix_styled.shape[1]
                    ws_summary.merge_cells(start_row=1, start_column=1, end_row=1, end_column=num_cols)
                    title_cell = ws_summary.cell(row=1, column=1, value=report_month_year)
                    title_cell.font = font_title
                    title_cell.fill = green_fill
                    title_cell.alignment = Alignment(horizontal="center", vertical="center")
                    ws_summary.row_dimensions[1].height = 35
                    
                    for c_idx in range(1, num_cols + 1):
                        ws_summary.cell(row=1, column=c_idx).fill = green_fill
                        ws_summary.cell(row=1, column=c_idx).border = thin_border
                    
                    ws_summary.row_dimensions[2].height = 24
                    for c_idx, col_name in enumerate(matrix_styled.columns, start=1):
                        cell = ws_summary.cell(row=2, column=c_idx, value=col_name)
                        cell.font = font_header
                        cell.fill = mint_header_fill
                        cell.border = thin_border
                        
                        # label center
                        if col_name == 'RFD':
                            cell.alignment = Alignment(horizontal="center", vertical="center")
                        else:
                            cell.alignment = Alignment(horizontal="center", vertical="center")
                            
                    for r_idx, row_data in enumerate(dataframe_to_rows(matrix_styled, index=False, header=False), start=3):
                        ws_summary.row_dimensions[r_idx].height = 20
                        is_grand_total_row = (row_data[0] == "Grand Total")
                        
                        row_fill = zebra_fill if (r_idx % 2 == 0 and not is_grand_total_row) else None
                        
                        for c_idx, val in enumerate(row_data, start=1):
                            display_val = "" if (isinstance(val, (int, float)) and val == 0 and c_idx != 1) else val
                            
                            cell = ws_summary.cell(row=r_idx, column=c_idx, value=display_val)
                            cell.border = thin_border
                            if row_fill:
                                cell.fill = row_fill
                                
                            if is_grand_total_row:
                                cell.font = font_body_bold
                                if c_idx == 1:
                                    cell.value = "TOTAL" # Match image terminology
                                    cell.alignment = Alignment(horizontal="center", vertical="center")
                                else:
                                    cell.alignment = Alignment(horizontal="center", vertical="center")
                            else:
                                if c_idx == 1:
                                    cell.font = font_body_regular
                                    cell.alignment = Alignment(horizontal="center", vertical="center")
                                else:
                                    cell.font = font_body_regular
                                    cell.alignment = Alignment(horizontal="center", vertical="center")
                                    
                    # autofit
                    for col in ws_summary.columns:
                        max_len = 0
                        col_letter = openpyxl.utils.get_column_letter(col[0].column)
                        
                        for cell in col:
                            if cell.row == 1:
                                continue
                            if cell.value:
                                max_len = max(max_len, len(str(cell.value)))
                                
                        ws_summary.column_dimensions[col_letter].width = max(max_len + 5, 12)
                        
                    
                    wb_styled.save(output_buffer)
                else:
                    with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
                        df_final_cleaned.to_excel(writer, index=False, sheet_name='DAR')
                        matrix.to_excel(writer, index=False, sheet_name='RFD Summary Matrix')
                    
                month_year_suffix = generate_month_year_suffix(df_final_cleaned)
                summary_filename = f"RFD Summary {month_year_suffix}.xlsx"

                st.subheader("Your file is ready.")
                st.download_button(
                    label=f"Download {summary_filename}",
                    data=output_buffer.getvalue(),
                    file_name=summary_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
        # pls pls pls wag na pls Lord
        except Exception as e:
            st.error(f"An error occurred while compiling your dynamic summary data layout: {e}")

# COD AUTOMATION PAGE ===============================================================================================================================================================
elif st.session_state.current_page == "cod_automation":

    WORD_TEMPLATE = Path(r"\\192.168.15.241\admin\ACTIVE\hpdaludado\SBC INSURANCE\COD\TEMPLATE\SBCI COD TEMPLATE.docx")
    EXCEL_TEMPLATE = Path(r"\\192.168.15.241\admin\ACTIVE\hpdaludado\SBC INSURANCE\COD\TEMPLATE\SBCI COD TEMPLATE.xlsx")

    st.markdown(
        """
        <div style="text-align: center; line-height: 1.2;">
            <h1 style="margin: 0; padding: 0;">SBC INSURANCE</h1>
            <p style="margin: 0; padding: 0; font-size: 35px;">
                Certificate of Destruction
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # =========================
    # CUSTOM INPUTS (WITH DEFAULTS + DROPDOWNS)
    # =========================
    col1, col2 = st.columns(2)

    with col1:
        custom_date = st.date_input("📅 Certificate Date", value=date.today())
        prepared_date = st.date_input("📝 Prepared Date", value=date.today())

    with col2:
        months_2026 = [
            "January 2026", "February 2026", "March 2026",
            "April 2026", "May 2026", "June 2026",
            "July 2026", "August 2026", "September 2026",
            "October 2026", "November 2026", "December 2026"
        ]

        custom_month = st.selectbox("📆 Month", months_2026)

        custom_reason = st.selectbox(
            "📌 Reason",
            ["Ceased/Paid", "End of Retention"]
        )

    pull_out_date = st.date_input("📤 Pull-Out Date", value=date.today())

    # =========================
    # SAFE TEXT REPLACEMENT
    # =========================
    def replace_text(doc, replacements, bold_keys=None):
        bold_keys = bold_keys or []

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:

                    if key in bold_keys:
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)
                                run.bold = True
                    else:
                        paragraph.text = paragraph.text.replace(key, value)

    # =========================
    # TABLE BORDERS
    # =========================
    def set_table_borders(table):
        tbl = table._element
        tblPr = tbl.tblPr

        borders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)

        tblPr.append(borders)

    # =========================
    # MAIN TABLE
    # =========================
    def insert_table(doc, df):
        for paragraph in doc.paragraphs:
            if "{{TABLE}}" in paragraph.text:

                paragraph.text = ""

                table = doc.add_table(rows=1, cols=len(df.columns))
                table.autofit = True
                set_table_borders(table)

                for col_index, col_name in enumerate(df.columns):
                    cell = table.rows[0].cells[col_index]
                    cell.text = str(col_name)

                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True

                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), "4F81BD")
                    tcPr.append(shd)

                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for col_index, value in enumerate(row):
                        cell = row_cells[col_index]
                        cell.text = str(value) if pd.notna(value) else ""
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                paragraph._element.addnext(table._element)
                break

    # =========================
    # SUMMARY GENERATION
    # =========================
    def generate_summary(df):
        if "BUCKET LEVEL" not in df.columns or "BALANCE" not in df.columns:
            return None

        df["BALANCE"] = pd.to_numeric(df["BALANCE"], errors="coerce").fillna(0)

        summary = df.groupby("BUCKET LEVEL").agg(
            TOTAL_ACCTS=("BUCKET LEVEL", "count"),
            TOTAL_BALANCE=("BALANCE", "sum")
        ).reset_index()

        total_accts = summary["TOTAL_ACCTS"].sum()
        total_balance = summary["TOTAL_BALANCE"].sum()

        return summary, total_accts, total_balance

    # =========================
    # SUMMARY TABLE
    # =========================
    def insert_summary_table(doc, summary, total_accts, total_balance):
        for paragraph in doc.paragraphs:
            if "{{SUMMARY_TABLE}}" in paragraph.text:

                paragraph.text = ""

                table = doc.add_table(rows=1, cols=3)
                table.autofit = True
                set_table_borders(table)

                summary_row = table.rows[0].cells
                merged_cell = summary_row[0].merge(summary_row[1]).merge(summary_row[2])
                merged_cell.text = "SUMMARY"

                for p in merged_cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True

                tc = merged_cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), "4F81BD")
                tcPr.append(shd)

                headers = ["Bucket Level", "Total # of Accts", "Total Balance"]
                header_cells = table.add_row().cells

                for i, header in enumerate(headers):
                    cell = header_cells[i]
                    cell.text = header

                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True

                for _, row in summary.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row["BUCKET LEVEL"])
                    row_cells[1].text = str(row["TOTAL_ACCTS"])
                    row_cells[2].text = f"{row['TOTAL_BALANCE']:,.2f}"

                    for cell in row_cells:
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                total_row = table.add_row().cells
                total_row[0].text = "Grand Total"
                total_row[1].text = str(total_accts)
                total_row[2].text = f"{total_balance:,.2f}"

                for cell in total_row:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True

                paragraph._element.addnext(table._element)
                break

    # =========================
    # PRODUCT + BUCKET TEXT
    # =========================
    def generate_product_bucket_text(df):
        if "PRODUCT" not in df.columns or "BUCKET LEVEL" not in df.columns:
            return ""

        products = sorted(df["PRODUCT"].dropna().astype(str).str.strip().unique())
        buckets = sorted(df["BUCKET LEVEL"].dropna().astype(str).str.strip().unique())

        return f"<{', '.join(products)} / {', '.join(buckets)}>"

    # =========================
    # MAIN PROCESS
    # =========================
    if st.button("🚀 Generate Certificate"):

        try:
            df = pd.read_excel(EXCEL_TEMPLATE)

            df.columns = df.columns.str.strip().str.upper()
            original_df = df.copy()

            rename_map = {}
            for col in df.columns:
                if "IS" in col:
                    rename_map[col] = "IS#"
                if "ENDO" in col:
                    rename_map[col] = "ENDO DATE"
                if "PULL" in col:
                    rename_map[col] = "PULL-OUT DATE"

            df = df.rename(columns=rename_map)

            if "NO" not in df.columns:
                df.insert(0, "NO", range(1, len(df) + 1))

            formatted_date = custom_date.strftime("%B %d, %Y")
            formatted_prepared_date = prepared_date.strftime("%m/%d/%Y")
            formatted_pullout_date = pull_out_date.strftime("%m/%d/%Y")

            df["PULL-OUT DATE"] = formatted_pullout_date

            if "ENDO DATE" in df.columns:
                df["ENDO DATE"] = pd.to_datetime(df["ENDO DATE"], errors="coerce").dt.strftime("%m/%d/%Y")

            df = df[["NO", "IS#", "ENDO DATE", "PULL-OUT DATE"]]

            st.success("✅ Excel Loaded Successfully!")
            st.dataframe(df)

            product_bucket_text = generate_product_bucket_text(original_df)

            doc = Document(WORD_TEMPLATE)

            replacements = {
                "{{DATE}}": formatted_date,
                "{{MONTH}}": custom_month,
                "{{REASON}}": custom_reason,
                "{{PREPARED_DATE}}": formatted_prepared_date,
                "{{PRODUCT_BUCKET}}": product_bucket_text
            }

            replace_text(
                doc,
                replacements,
                bold_keys=["{{MONTH}}", "{{DATE}}"]
            )

            insert_table(doc, df)

            summary_data = generate_summary(original_df)

            if summary_data:
                summary_df, total_accts, total_balance = summary_data
                insert_summary_table(doc, summary_df, total_accts, total_balance)

            OUTPUT_FOLDER = Path(r"\\192.168.15.241\admin\ACTIVE\hpdaludado\SBC INSURANCE\COD")
            OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)

            # Use the selected Certificate Date
            certificate_date = custom_date.strftime("%m-%d-%Y")

            # Filename based on selected reason
            if custom_reason == "End of Retention":
                filename = (
                    f"SBC INSURANCE_Certificate of Deletion and Destruction_"
                    f"END OF RETENTION_{certificate_date}.docx"
                )
            else:
                filename = (
                    f"SBC INSURANCE_Certificate of Deletion and Destruction_"
                    f"CEASED_PAID_{certificate_date}.docx"
                )

            output_file = OUTPUT_FOLDER / filename
            doc.save(output_file)

            st.success(f"✅ Saved successfully!\n\n{output_file}")

        except Exception as e:
            st.error(f"❌ Error: {e}")